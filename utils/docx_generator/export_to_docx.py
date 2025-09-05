from docx import Document
import os
import json


def export_response_to_docx(
    docx_filename,
    result_heading,
    request_url,
    request_payload,
    request_headers,
    response_text,
    response_headers,
    append=False
):
    """Export API request/response details into a .docx file."""
    if append and os.path.exists(docx_filename):
        document = Document(docx_filename)

        document.add_paragraph(result_heading, style="Heading 1")

        document.add_paragraph("Request URL:", style="Heading 2")
        document.add_paragraph(request_url)

        if request_headers:
            document.add_paragraph("Request Headers:", style="Heading 2")
            for key, value in request_headers.items():
                document.add_paragraph(f"{key}: {value}", style="NoSpacing")

        if request_payload:
            document.add_paragraph("Request Payload (JSON):", style="Heading 2")
            document.add_paragraph(request_payload)

        if response_text:
            document.add_paragraph("Response Body:", style="Heading 2")
            document.add_paragraph(response_text)

        if response_headers:
            document.add_paragraph("Response Headers:", style="Heading 2")
            for key, value in response_headers.items():
                document.add_paragraph(f"{key}: {value}", style="NoSpacing")

        try:
            document.save(docx_filename)
        except PermissionError:
            print(f"Permission denied: Could not write to '{docx_filename}'")
    else:
        print(f"[WARN] File '{docx_filename}' does not exist. Run the base .docx generator first.")


def export_all_results_to_docx(test_results, output_dir):
    """Append only PASSED test results into their respective .docx files."""
    for result in test_results:
        if result.get("status") != "Passed":
            continue  # Skip Failed/Error cases

        existing_docx = None
        for fname in os.listdir(output_dir):
            if fname.startswith(f"{result['tc_no']}") and fname.endswith(".docx"):
                existing_docx = os.path.join(output_dir, fname)
                break

        if not existing_docx or not os.path.exists(existing_docx):
            print(f"[WARN] Base docx for Test Case {result['tc_no']} not found in {output_dir}")
            continue

        for req in result["requests"]:
            export_response_to_docx(
                existing_docx,
                req["result_heading"],
                req["url"],
                json.dumps(req["body"], indent=2) if req["body"] else None,
                req["headers"],
                json.dumps(req["response_body"], indent=2),
                req["response_headers"],
                append=True
            )
